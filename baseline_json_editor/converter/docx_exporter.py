import io
import logging
from typing import Optional
from converter.schema import BaselineDocument

logger = logging.getLogger(__name__)

def export_edited_docx(baseline_doc: BaselineDocument) -> Optional[bytes]:
    """
    Converts a BaselineDocument schema into a fully formatted Microsoft Word (.docx) file.
    
    Args:
        baseline_doc: The edited BaselineDocument schema
        
    Returns:
        The new DOCX as bytes, or None if it fails.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title of document
        if baseline_doc.title:
            doc.add_heading(baseline_doc.title, 0)
            
        def process_blocks(blocks):
            for block in blocks:
                if not block.content and not block.children:
                    continue
                    
                content = block.content or ""
                is_blockquote = block.properties and block.properties.blockquote
                
                if block.block_type == "SectionHeader":
                    level = 1 # docx heading levels are 1-9
                    if block.properties and block.properties.heading_level:
                        level = min(9, max(1, block.properties.heading_level))
                    doc.add_heading(content, level=level)
                
                elif block.block_type == "Text":
                    p = doc.add_paragraph(content)
                    if is_blockquote:
                        p.style = 'Quote'
                
                elif block.block_type == "List":
                    style = 'List Number' if (block.properties and block.properties.list_type == "ordered") else 'List Bullet'
                    if content:
                        doc.add_paragraph(content) # optional leading text
                    if block.children:
                        for child in block.children:
                            child_content = child.content or ""
                            doc.add_paragraph(child_content, style=style)
                    continue # handled children
                
                elif block.block_type == "Table":
                    # Simplified table logic
                    if block.children:
                        # Assumes format: Table -> [Row, Row] -> [Cell, Cell]
                        rows = len(block.children)
                        cols = len(block.children[0].children) if rows > 0 and block.children[0].children else 1
                        
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for r_idx, row in enumerate(block.children):
                            for c_idx, cell in enumerate(row.children):
                                if c_idx < cols:
                                    cell_content = cell.content or ""
                                    table.cell(r_idx, c_idx).text = cell_content
                    else:
                        p = doc.add_paragraph(content)
                        p.style = 'Macro Text'
                
                elif block.block_type == "Code":
                    p = doc.add_paragraph(content)
                    p.style = 'Macro Text'
                
                elif block.block_type == "Equation":
                    p = doc.add_paragraph(f"[Equation] {content}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                else:
                    doc.add_paragraph(content)
                
                if block.children and block.block_type not in ["List", "Table"]:
                    process_blocks(block.children)

        for page in baseline_doc.pages:
            process_blocks(page.blocks)
            if page != baseline_doc.pages[-1]:
                doc.add_page_break()

        # Save to memory stream
        target_stream = io.BytesIO()
        doc.save(target_stream)
        
        return target_stream.getvalue()
        
    except ImportError:
        logger.error("python-docx is not installed. Cannot export DOCX.")
        return None
    except Exception as e:
        logger.error(f"Error during DOCX export: {e}")
        return None
