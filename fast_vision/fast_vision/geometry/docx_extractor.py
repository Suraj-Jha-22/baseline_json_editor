"""
Extract structured blocks from DOCX files using python-docx (MIT license).

Walks the document body to extract paragraphs, tables, and styles into
the same block-dict format the geometry layer produces, so the rest of
the pipeline (style normalizer, Vision API tagger, schema assembler) works
identically for both PDF and DOCX inputs.
"""

from __future__ import annotations

import logging
import uuid
from typing import Any, Dict, List, Tuple

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu

logger = logging.getLogger(__name__)

# We treat DOCX as a single-page document unless section breaks are present.
# Default page in points (US Letter 8.5×11 inches)
DEFAULT_WIDTH = 612.0
DEFAULT_HEIGHT = 792.0


def extract_from_docx(docx_path: str) -> Tuple[
    List[Dict[str, Any]],           # pages_data  (same shape as char_extractor output)
    Dict[int, List[Dict[str, Any]]],  # blocks_by_page
    Dict[int, List[Dict[str, Any]]],  # tables_by_page
]:
    """Extract paragraphs and tables from a DOCX file.

    Returns the same triple that the PDF geometry pipeline produces so the
    downstream pipeline can treat them identically.
    """
    doc = DocxDocument(docx_path)

    # ── Page dimensions from section ─────────────────────────────────
    page_w, page_h = DEFAULT_WIDTH, DEFAULT_HEIGHT
    if doc.sections:
        sec = doc.sections[0]
        if sec.page_width:
            page_w = _emu_to_pt(sec.page_width)
        if sec.page_height:
            page_h = _emu_to_pt(sec.page_height)

    # ── Walk body elements in document order ─────────────────────────
    blocks: List[Dict[str, Any]] = []
    tables_out: List[Dict[str, Any]] = []

    # Track vertical cursor for synthetic bbox assignment
    y_cursor = 36.0  # start ~0.5 inch from top margin
    left_margin = 72.0  # ~1 inch
    text_width = page_w - 2 * left_margin

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue

            text = _full_paragraph_text(para).strip()
            if not text:
                y_cursor += 6  # empty line spacing
                continue

            # Extract font info from first run (including runs inside hyperlinks)
            font_name = "Calibri"
            font_size = 11.0
            font_color = "#000000"
            is_bold = False
            is_italic = False
            alignment = "left"

            first_run = _first_run(para)
            if first_run:
                if first_run.font.name:
                    font_name = first_run.font.name
                if first_run.font.size:
                    font_size = first_run.font.size.pt
                if first_run.font.color and first_run.font.color.rgb:
                    font_color = f"#{first_run.font.color.rgb}"
                is_bold = bool(first_run.bold)
                is_italic = bool(first_run.italic)

            if para.alignment is not None:
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                }
                alignment = align_map.get(para.alignment, "left")

            # Build fontname string matching PDF convention
            fn = font_name
            if is_bold:
                fn += "-Bold"
            if is_italic:
                fn += "-Italic"

            # Estimate block height from font size + line count
            line_height = font_size * 1.4
            num_lines = max(1, len(text) * font_size * 0.6 / text_width + 0.5)
            block_height = line_height * num_lines

            # Synthesize a bbox
            bbox_x0 = left_margin
            bbox_y0 = y_cursor
            bbox_x1 = left_margin + text_width
            bbox_y1 = y_cursor + block_height

            # Build word-level tokens
            words = []
            word_x = bbox_x0
            for word_text in text.split():
                word_w = len(word_text) * font_size * 0.55
                words.append({
                    "text": word_text,
                    "x0": round(word_x, 2),
                    "y0": round(bbox_y0, 2),
                    "x1": round(min(word_x + word_w, bbox_x1), 2),
                    "y1": round(bbox_y1, 2),
                    "fontname": fn,
                    "size": font_size,
                    "color": font_color,
                })
                word_x += word_w + font_size * 0.3  # space
                if word_x > bbox_x1:
                    word_x = bbox_x0  # wrap

            # Guess block type from style
            style_name = (para.style.name or "").lower() if para.style else ""
            block_type, role = _classify_paragraph(style_name, font_size, is_bold, text)

            blocks.append({
                "id": str(uuid.uuid4()),
                "text": text,
                "x0": round(bbox_x0, 2),
                "y0": round(bbox_y0, 2),
                "x1": round(bbox_x1, 2),
                "y1": round(bbox_y1, 2),
                "fontname": fn,
                "size": font_size,
                "color": font_color,
                "alignment": alignment,
                "words": words,
                "block_type": block_type,
                "role": role,
                "reading_order": len(blocks),
            })

            y_cursor = bbox_y1 + font_size * 0.4  # paragraph spacing

        elif tag == "tbl":
            tbl = _find_table(doc, element)
            if tbl is None:
                continue

            rows_data = []
            for row in tbl.rows:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text.strip())
                rows_data.append(row_cells)

            n_rows = len(rows_data)
            n_cols = max(len(r) for r in rows_data) if rows_data else 0

            if n_rows == 0 or n_cols == 0:
                continue

            table_height = n_rows * 20.0
            tbl_bbox = [
                round(left_margin, 2),
                round(y_cursor, 2),
                round(left_margin + text_width, 2),
                round(y_cursor + table_height, 2),
            ]

            cells = []
            col_w = text_width / max(n_cols, 1)
            row_h = table_height / max(n_rows, 1)
            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row):
                    cells.append({
                        "row": r_idx,
                        "col": c_idx,
                        "row_span": 1,
                        "col_span": 1,
                        "text": cell_text,
                        "bbox": [
                            round(left_margin + c_idx * col_w, 2),
                            round(y_cursor + r_idx * row_h, 2),
                            round(left_margin + (c_idx + 1) * col_w, 2),
                            round(y_cursor + (r_idx + 1) * row_h, 2),
                        ],
                    })

            tables_out.append({
                "id": str(uuid.uuid4()),
                "page": 1,
                "rows": n_rows,
                "cols": n_cols,
                "bbox": tbl_bbox,
                "cells": cells,
                "block_type": "table",
            })

            y_cursor += table_height + 12.0

    # If document is long, paginate at page_h boundaries
    pages_data, final_blocks, final_tables = _paginate(
        blocks, tables_out, page_w, page_h,
    )

    logger.info(
        "DOCX extracted: %d blocks, %d tables, %d pages",
        len(blocks), len(tables_out), len(pages_data),
    )

    return pages_data, final_blocks, final_tables


# ── Helpers ──────────────────────────────────────────────────────────

def _emu_to_pt(emu_val) -> float:
    """Convert EMU (English Metric Units) to points."""
    if hasattr(emu_val, "pt"):
        return float(emu_val.pt)
    return float(emu_val) / 12700.0


def _find_paragraph(doc, element):
    """Find the python-docx Paragraph object for a given XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Find the python-docx Table object for a given XML element."""
    for tbl in doc.tables:
        if tbl._element is element:
            return tbl
    return None


def _full_paragraph_text(para) -> str:
    """Extract ALL text from a paragraph, including hyperlinks.

    python-docx's para.text only yields text from direct <w:r> children,
    but hyperlinks (<w:hyperlink>) contain their own <w:r> children that
    are silently skipped.  This function walks every child element to
    capture everything.
    """
    from lxml import etree

    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts = []

    for child in para._element:
        tag = etree.QName(child).localname

        if tag == "r":
            # Direct run
            for t_el in child.findall("w:t", nsmap):
                if t_el.text:
                    parts.append(t_el.text)

        elif tag == "hyperlink":
            # Runs inside a hyperlink — this is what para.text misses
            for run_el in child.findall("w:r", nsmap):
                for t_el in run_el.findall("w:t", nsmap):
                    if t_el.text:
                        parts.append(t_el.text)

    return "".join(parts)


def _first_run(para):
    """Get the first Run from a paragraph, even if it's inside a hyperlink.

    para.runs only returns direct <w:r> children.  If the paragraph
    starts with a hyperlink, para.runs may be empty while there are
    perfectly good runs inside the hyperlink.
    """
    from docx.text.run import Run
    from lxml import etree

    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for child in para._element:
        tag = etree.QName(child).localname

        if tag == "r":
            return Run(child, para)

        if tag == "hyperlink":
            for run_el in child.findall("w:r", nsmap):
                return Run(run_el, para)

    return None


def _classify_paragraph(
    style_name: str, font_size: float, is_bold: bool, text: str
) -> Tuple[str, str]:
    """Classify a paragraph into block_type and role from DOCX style info."""
    # Heading styles
    if "heading" in style_name or "title" in style_name:
        if "title" in style_name:
            return "heading", "title"
        # heading 1, heading 2, etc.
        for lvl in range(1, 7):
            if str(lvl) in style_name:
                if lvl <= 2:
                    return "heading", "section_title"
                return "heading", "subsection_title"
        return "heading", "section_title"

    # List styles
    if "list" in style_name or "bullet" in style_name:
        return "list_item", "list_item"

    # Caption
    if "caption" in style_name:
        return "caption", "caption"

    # Code / monospace
    if "code" in style_name or "mono" in style_name:
        return "code_block", "paragraph"

    # Heuristic: large bold text = heading
    if is_bold and font_size >= 14:
        return "heading", "section_title"

    # Bullet patterns in text
    stripped = text.lstrip()
    if stripped.startswith(("•", "–", "—", "▪", "◦", "○")):
        return "list_item", "list_item"
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)":
        return "list_item", "list_item"

    return "paragraph", "paragraph"


def _paginate(
    blocks: List[Dict[str, Any]],
    tables: List[Dict[str, Any]],
    page_w: float,
    page_h: float,
) -> Tuple[
    List[Dict[str, Any]],            # pages_data
    Dict[int, List[Dict[str, Any]]],  # blocks_by_page
    Dict[int, List[Dict[str, Any]]],  # tables_by_page
]:
    """Split blocks across pages based on y-coordinates."""
    pages_data: List[Dict[str, Any]] = []
    blocks_by_page: Dict[int, List[Dict[str, Any]]] = {}
    tables_by_page: Dict[int, List[Dict[str, Any]]] = {}

    if not blocks and not tables:
        pages_data.append({
            "page_number": 1,
            "width": page_w,
            "height": page_h,
            "chars": [],
        })
        blocks_by_page[0] = []
        tables_by_page[0] = []
        return pages_data, blocks_by_page, tables_by_page

    # Determine how many pages we need
    max_y = max(
        max((b["y1"] for b in blocks), default=0),
        max((t["bbox"][3] for t in tables), default=0),
    )
    n_pages = max(1, int(max_y / page_h) + 1)

    for page_num in range(1, n_pages + 1):
        pages_data.append({
            "page_number": page_num,
            "width": page_w,
            "height": page_h,
            "chars": [],  # DOCX doesn't have raw chars
        })
        page_idx = page_num - 1
        page_top = page_idx * page_h
        page_bottom = page_top + page_h

        page_blocks = []
        for b in blocks:
            if b["y0"] >= page_top and b["y0"] < page_bottom:
                # Adjust coordinates relative to this page
                adjusted = dict(b)
                adjusted["y0"] = b["y0"] - page_top
                adjusted["y1"] = b["y1"] - page_top
                # Adjust word bboxes too
                if "words" in adjusted:
                    adjusted["words"] = [
                        {**w, "y0": w["y0"] - page_top, "y1": w["y1"] - page_top}
                        for w in adjusted["words"]
                    ]
                page_blocks.append(adjusted)

        blocks_by_page[page_idx] = page_blocks

        page_tables = []
        for t in tables:
            if t["bbox"][1] >= page_top and t["bbox"][1] < page_bottom:
                adjusted = dict(t)
                adjusted["bbox"] = [
                    t["bbox"][0],
                    t["bbox"][1] - page_top,
                    t["bbox"][2],
                    t["bbox"][3] - page_top,
                ]
                adjusted["cells"] = [
                    {**c, "bbox": [c["bbox"][0], c["bbox"][1] - page_top,
                                   c["bbox"][2], c["bbox"][3] - page_top]}
                    for c in t.get("cells", [])
                ]
                page_tables.append(adjusted)

        tables_by_page[page_idx] = page_tables

    return pages_data, blocks_by_page, tables_by_page
